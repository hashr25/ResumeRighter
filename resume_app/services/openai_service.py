import openai
from django.conf import settings
import fitz  # PyMuPDF
import docx
from docx import Document
from io import BytesIO
from typing import Union, Dict

# Initialize OpenAI API key
openai.api_key = settings.OPENAI_API_KEY
openai.log = "debug"

openai_model = settings.OPENAI_MODEL


# Service functions
def validate_resume(file: BytesIO, file_type: str) -> Dict[str, Union[bool, str]]:
    """
    Validate that the uploaded file content looks like a resume.
    :param file: The file object (BytesIO for uploaded files).
    :param file_type: The type of the file ('pdf', 'docx', 'txt').
    :return: A dictionary with is_valid (bool) and validated_data (str) if valid.
    """
    try:
        # Step 1: Extract text from the file
        file_content = extract_text_from_file(file, file_type)
        if not file_content:
            return {"is_valid": False, "validated_data": ""}

        # Step 2: Validate the extracted text
        response = openai.chat.completions.create(
            model=openai_model,
            messages=[
                {"role": "system", "content": "You are an AI assistant for validating resumes."},
                {"role": "user",
                 "content": f"Does the following text look like a person's resume?\n\n{file_content}\n\nRespond 'yes' or 'no'."},
            ],
            temperature=0,
            max_tokens=10,
        )

        result = response.choices[0].message.content.strip().lower()
        return {"is_valid": result == "yes", "validated_data": file_content if result == "yes" else ""}
    except openai.OpenAIError as e:
        print(f"OpenAI API error: {e}")
        return {"is_valid": False, "validated_data": ""}
    except Exception as e:
        print(f"Unexpected error: {e}")
        return {"is_valid": False, "validated_data": ""}


def validate_job_posting(url: str) -> Dict[str, Union[bool, str]]:
    """
    Validate that the given URL points to a job posting.
    :param url: The job posting URL.
    :return: A dictionary with is_valid (bool) and validated_data (str) if valid.
    """
    try:
        response = openai.chat.completions.create(
            model=openai_model,
            messages=[
                {"role": "system", "content": "You are an AI assistant for validating job postings."},
                {"role": "user",
                 "content": f"Does the following URL point to a job posting?\n\n{url}\n\nRespond 'yes' or 'no'."},
            ],
            temperature=0,
            max_tokens=10,
        )
        result = response.choices[0].message.content.strip().lower()
        return {"is_valid": result == "yes", "validated_data": url if result == "yes" else ""}
    except openai.OpenAIError as e:
        print(f"OpenAI API error: {e}")
        return {"is_valid": False, "validated_data": ""}
    except Exception as e:
        print(f"Unexpected error: {e}")
        return {"is_valid": False, "validated_data": ""}


def validate_special_considerations(text: str) -> Dict[str, Union[bool, str]]:
    """
    Validate the special considerations text provided by the user.
    :param text: The special considerations input.
    :return: A dictionary with is_valid (bool) and validated_data (str) if valid.
    """
    try:
        response = openai.chat.completions.create(
            model=openai_model,
            messages=[
                {"role": "system", "content": "You are an AI assistant for validating special considerations."},
                {"role": "user",
                 "content": f"Is the following text relevant to resume optimization?\n\n{text}\n\nRespond 'yes' or 'no'."},
            ],
            temperature=0,
            max_tokens=10,
        )
        result = response.choices[0].message.content.strip().lower()
        return {"is_valid": result == "yes", "validated_data": text if result == "yes" else ""}
    except openai.OpenAIError as e:
        print(f"OpenAI API error: {e}")
        return {"is_valid": False, "validated_data": ""}
    except Exception as e:
        print(f"Unexpected error: {e}")
        return {"is_valid": False, "validated_data": ""}


def extract_text_from_file(file: Union[BytesIO, str], file_type: str) -> str:
    """
    Extract text content from a file based on its type.
    :param file: The file object (BytesIO for uploaded files or path for local files).
    :param file_type: The type of the file ('pdf', 'docx', 'txt').
    :return: Extracted text as a string.
    """
    try:
        if file_type == "pdf":
            text = ""
            pdf_document = fitz.open(stream=file.read(), filetype="pdf")
            for page in pdf_document:
                text += page.get_text()
            pdf_document.close()
            return text
        elif file_type == "docx":
            doc = docx.Document(file)
            return "\n".join([p.text for p in doc.paragraphs])
        elif file_type == "txt":
            return file.read().decode("utf-8")
        else:
            raise ValueError("Unsupported file type.")
    except Exception as e:
        print(f"Error extracting text from {file_type}: {e}")
        return ""


def generate_rewritten_resume(
        resume_text: str, job_posting_text: str, considerations: str
) -> bytes:
    """
    Generate a rewritten resume based on the original resume, job posting, and special considerations.
    """
    extra_details = settings.EXTRA_DETAILS_FOR_RESUME_GENERATION

    try:
        prompt = (
            "Your goal is to make suggestions for each job so it aligns better with the provided job posting, highlights relevant skills and experience, "
            "and adheres to professional standards. Follow these instructions:\n\n"
            f"{extra_details}\n\n"
            "Specific Guidelines:\n"
            "1. Make suggestions for all jobs listed in the original resume. Do not remove any job entries.\n"
            "2. Rewrite descriptions and bullet points for each job posting to better align with the job posting. Use action-oriented language.\n"
            "3. Highlight transferable skills and technologies relevant to the job posting. Add keywords from the job posting where appropriate.\n"
            "4. Limit the suggestions for each job listing to no more than 5 bullet points.\n\n"
            "5. Do not list individual skills or technologies separately. Instead, incorporate them into the job descriptions.\n\n"
            "6. Do not include personal information, such as name, address, or contact details, in the resume suggestion document.\n\n"
            "7. Do not include any education or certification information in the resume suggestion document.\n\n"
            "Original Resume:\n\n"
            f"{resume_text}\n\n"
            "Job Posting:\n\n"
            f"{job_posting_text}\n\n"
            "Special Considerations (if provided):\n\n"
            f"{considerations if considerations.strip() else 'None'}\n\n"
            "Output the suggestions for the resume in a clean format, labeling each job by company name.\n\n"
            "Include a footer only if the name \"Randy Hash\" is not present in the resume. The footer should say:\n"
            "'This resume was generated using Resume Righter, written by Randy Hash. The source code for this project can be found at {insert link to https://github.com/hashr25/ResumeRighter that simple says 'GitHub'}'."
        )

        # Generate rewritten resume using OpenAI API
        response = openai.chat.completions.create(
            model=openai_model,
            messages=[
                {"role": "system", "content": "You are a professional resume consultant tasked with improving a resume by rewriting bullet points for previous jobs. "},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=4096,
        )
        rewritten_text = response.choices[0].message.content.strip()

        # Create a .docx file with the rewritten resume
        document = Document()
        for paragraph in rewritten_text.split("\n\n"):
            document.add_paragraph(paragraph.strip())

        # Save the document to a byte stream
        byte_stream = BytesIO()
        document.save(byte_stream)
        byte_stream.seek(0)
        return byte_stream.getvalue()

    except openai.OpenAIError as e:
        print(f"OpenAI API error: {e}")
        raise
    except Exception as e:
        print(f"Unexpected error: {e}")
        raise
